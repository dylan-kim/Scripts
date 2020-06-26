# -*- coding: utf-8 -*-

# Python 3.7
# python-docx 0.8.10

import os
import datetime
import locale
import pathlib
import shutil

from docx import Document
from docx.text.run import Font
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.shared import Pt


locale.setlocale(locale.LC_ALL, 'fr_FR')

name = "Dylan KIM"
subject = "Composants logiciels"
engineer_year = "I3"
option = "LD"

date = datetime.datetime.now()
file_name = "Rapport TP"

def create_report():
    document = Document()

    # Header
    section = document.sections[0]
    header = section.header
    paragraph  = header.paragraphs[0]
    month_year = date.strftime("%B").upper() + " " + str(date.year)
    paragraph.text = name + " " + engineer_year + " " + option + "\t\t" + month_year # \t is a separator between left/center/right
    paragraph.style = document.styles["Header"]

    # Title
    first_paragraph = document.add_paragraph()
    first_paragraph.alignment  = WD_ALIGN_PARAGRAPH.CENTER
    first_paragraph = first_paragraph.add_run(file_name)

    font = first_paragraph.font
    font.name = 'Calibri (Corps)'
    font.size = Pt(19)
    font.color.rgb = RGBColor(24, 85, 178)

    # Body 
    document.add_heading('Introduction', level=1)
    document.add_heading('Conclusion', level=1)
    document.add_heading('Annexe', level=1)

    document.save(file_name + ".docx")

    # We cannot add page number with the actual library (footnote), so we can't automate that part.

def main():
    create_report()

    # Move the document to the desktop
    folder_destination = pathlib.Path(r'C:\Users\utilisateur\Desktop')
    shutil.copy2(file_name + ".docx", folder_destination)

    os.startfile(r'C:\Users\utilisateur\Desktop' + file_name + '.docx')

if __name__ == "__main__": 
    main()